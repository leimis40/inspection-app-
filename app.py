import os
import io
import json
import uuid
import sqlite3
from pathlib import Path
from contextlib import closing
from datetime import date, datetime
from typing import List, Dict, Any
from difflib import SequenceMatcher

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas

DB_PATH = os.getenv("INSPECTION_DB_PATH", "inspection_reports.db")
UPLOAD_DIR = Path(os.getenv("INSPECTION_UPLOAD_DIR", "uploads"))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
DATA_DIR = Path(os.getenv("INSPECTION_DATA_DIR", "data"))
DRAFT_DIR = Path(os.getenv("INSPECTION_DRAFT_DIR", "drafts"))
DRAFT_DIR.mkdir(parents=True, exist_ok=True)
CSV_KB_PATH = Path(os.getenv("INSPECTION_KB_CSV", DATA_DIR / "inspection_ai_knowledge_base.csv"))
FINDINGS_CSV_PATH = Path(os.getenv("INSPECTION_FINDINGS_CSV", DATA_DIR / "final_inspection_findings.csv"))
LEGAL_CSV_PATH = Path(os.getenv("INSPECTION_LEGAL_CSV", DATA_DIR / "final_legal_references.csv"))
STANDARDS_CSV_PATH = Path(os.getenv("INSPECTION_STANDARDS_CSV", DATA_DIR / "final_standards_clauses.csv"))

SCHEMA_SQL = """
PRAGMA foreign_keys = ON;

CREATE TABLE IF NOT EXISTS reports (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    report_type TEXT NOT NULL,
    report_number TEXT NOT NULL,
    report_date TEXT NOT NULL,
    place TEXT,

    site_name TEXT,
    site_address TEXT,
    equipment_class TEXT,
    main_fuse_size TEXT,
    grid_operator TEXT,
    site_description TEXT,

    holder_name TEXT,
    holder_address TEXT,
    holder_contact TEXT,
    operator_person TEXT,

    builder_name TEXT,
    builder_address TEXT,
    builder_contact TEXT,
    electrical_manager TEXT,

    commissioning_inspection TEXT,
    previous_periodic_inspection TEXT,
    this_inspection TEXT,
    next_periodic_inspection TEXT,
    participants TEXT,

    inspection_area TEXT,
    inspection_scope_note TEXT,
    inspected_part TEXT,
    method_text TEXT,
    norm_documents TEXT,

    decision_summary TEXT,
    immediate_danger_flag INTEGER DEFAULT 0,
    reinspection_required_flag INTEGER DEFAULT 0,
    reinspection_due TEXT,

    recommendations_text TEXT,
    attachments_text TEXT,
    distribution_text TEXT,
    signature_name TEXT,
    signature_title TEXT,

    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
    updated_at TEXT DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE IF NOT EXISTS findings (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    report_id INTEGER NOT NULL,
    severity_class TEXT NOT NULL,
    finding_no INTEGER NOT NULL,
    area TEXT,
    finding_text TEXT NOT NULL,
    standard_reference TEXT,
    legal_reference TEXT,
    corrective_action TEXT,
    status TEXT DEFAULT 'open',
    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (report_id) REFERENCES reports(id) ON DELETE CASCADE
);

CREATE TABLE IF NOT EXISTS finding_photos (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    finding_id INTEGER NOT NULL,
    original_filename TEXT NOT NULL,
    stored_filename TEXT NOT NULL,
    file_path TEXT NOT NULL,
    caption TEXT,
    created_at TEXT DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (finding_id) REFERENCES findings(id) ON DELETE CASCADE
);

CREATE TABLE IF NOT EXISTS measurements (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    report_id INTEGER NOT NULL,
    center_name TEXT,
    group_name TEXT,
    ik_a TEXT,
    zk_ohm TEXT,
    protection_a TEXT,
    phase_order TEXT,
    voltage_v TEXT,
    idn_ma TEXT,
    trip_time_ms TEXT,
    delta_current_ma TEXT,
    FOREIGN KEY (report_id) REFERENCES reports(id) ON DELETE CASCADE
);
"""

DEFAULT_METHOD_TEXT = (
    "Tarkastus on suoritettu noudattaen voimassa olevia sähköturvallisuussäädöksiä, "
    "hyvää hallintotapaa, tarkastajan omaa laatuohjeistusta ja soveltaen standardia SFS 5825."
)
DEFAULT_NORMS = "SFS 6000; SFS 6001; STL 1135/2016; VNa 1434/2016; VNa 576/2003"
DRAFT_FILE = DRAFT_DIR / "report_form_draft.json"


@st.cache_data(show_spinner=False)
def load_csv_if_exists(path_str: str) -> pd.DataFrame:
    path = Path(path_str)
    if path.exists() and path.is_file():
        try:
            return pd.read_csv(path)
        except Exception:
            return pd.DataFrame()
    return pd.DataFrame()


def _pick_first_existing(df: pd.DataFrame, candidates: List[str]) -> str:
    lower_map = {c.lower(): c for c in df.columns}
    for candidate in candidates:
        if candidate.lower() in lower_map:
            return lower_map[candidate.lower()]
    return ""


def _tokenize(text: str) -> List[str]:
    return [
        t.strip()
        for t in str(text).lower().replace("/", " ").replace(",", " ").replace(";", " ").split()
        if len(t.strip()) > 2
    ]


def suggest_references_from_csv(finding_text: str, limit: int = 5) -> List[dict]:
    suggestions: List[dict] = []
    query = str(finding_text or "").strip()
    if not query:
        return suggestions

    kb_df = load_csv_if_exists(str(CSV_KB_PATH))
    if kb_df.empty:
        standards_df = load_csv_if_exists(str(STANDARDS_CSV_PATH))
        legal_df = load_csv_if_exists(str(LEGAL_CSV_PATH))
        findings_df = load_csv_if_exists(str(FINDINGS_CSV_PATH))
        frames = []
        if not standards_df.empty:
            frames.append(standards_df.copy())
        if not legal_df.empty:
            frames.append(legal_df.copy())
        if not findings_df.empty:
            frames.append(findings_df.copy())
        if frames:
            kb_df = pd.concat(frames, ignore_index=True, sort=False)
        else:
            return suggestions

    text_col = _pick_first_existing(kb_df, ["full_text", "embedding_text", "summary", "text", "content", "finding_text"])
    ref_col = _pick_first_existing(kb_df, ["reference", "standard_reference", "legal_reference", "clause_number", "section", "pykala"])
    code_col = _pick_first_existing(kb_df, ["standard_code", "code", "law_code", "document_code", "source_name"])
    version_col = _pick_first_existing(kb_df, ["standard_version", "version", "year"])
    heading_col = _pick_first_existing(kb_df, ["heading", "title", "section_title"])
    type_col = _pick_first_existing(kb_df, ["chunk_type", "source_type", "document_type", "type"])

    if not text_col:
        return suggestions

    query_tokens = set(_tokenize(query))
    if not query_tokens:
        return suggestions

    rows = []
    for _, row in kb_df.fillna("").iterrows():
        hay = str(row.get(text_col, ""))
        if not hay:
            continue
        hay_tokens = set(_tokenize(hay))
        overlap = len(query_tokens & hay_tokens)
        if overlap == 0:
            continue
        similarity = SequenceMatcher(None, query.lower(), hay.lower()[:2000]).ratio()
        score = overlap * 10 + similarity * 5
        rows.append((score, row))

    rows.sort(key=lambda x: x[0], reverse=True)
    seen = set()
    for score, row in rows:
        code = str(row.get(code_col, "")).strip() if code_col else ""
        version = str(row.get(version_col, "")).strip() if version_col else ""
        ref = str(row.get(ref_col, "")).strip() if ref_col else ""
        heading = str(row.get(heading_col, "")).strip() if heading_col else ""
        source_type = str(row.get(type_col, "")).strip() if type_col else ""
        excerpt = str(row.get(text_col, "")).strip().replace("\n", " ")[:500]

        if code and version and ref and ref not in code:
            label = f"{code}:{version} {ref}"
        elif code and ref and ref not in code:
            label = f"{code} {ref}"
        elif code and version:
            label = f"{code}:{version}"
        else:
            label = ref or heading or "Viite"

        if label in seen:
            continue
        seen.add(label)
        suggestions.append(
            {
                "label": label,
                "heading": heading,
                "source_type": source_type,
                "excerpt": excerpt,
                "score": round(float(score), 2),
            }
        )
        if len(suggestions) >= limit:
            break
    return suggestions


def get_conn() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_db() -> None:
    with closing(get_conn()) as conn:
        conn.executescript(SCHEMA_SQL)
        conn.commit()


def save_uploaded_photo(uploaded_file, finding_id: int, caption: str, conn: sqlite3.Connection) -> None:
    ext = Path(uploaded_file.name).suffix.lower() or ".jpg"
    unique_name = f"finding_{finding_id}_{uuid.uuid4().hex}{ext}"
    target_path = UPLOAD_DIR / unique_name
    with open(target_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    conn.execute(
        """
        INSERT INTO finding_photos (finding_id, original_filename, stored_filename, file_path, caption)
        VALUES (?, ?, ?, ?, ?)
        """,
        (finding_id, uploaded_file.name, unique_name, str(target_path), caption),
    )
    conn.commit()


def create_report(data: dict, conn: sqlite3.Connection) -> int:
    cursor = conn.execute(
        """
        INSERT INTO reports (
            report_type, report_number, report_date, place,
            site_name, site_address, equipment_class, main_fuse_size, grid_operator, site_description,
            holder_name, holder_address, holder_contact, operator_person,
            builder_name, builder_address, builder_contact, electrical_manager,
            commissioning_inspection, previous_periodic_inspection, this_inspection, next_periodic_inspection, participants,
            inspection_area, inspection_scope_note, inspected_part, method_text, norm_documents,
            decision_summary, immediate_danger_flag, reinspection_required_flag, reinspection_due,
            recommendations_text, attachments_text, distribution_text, signature_name, signature_title
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            data["report_type"], data["report_number"], data["report_date"], data["place"],
            data["site_name"], data["site_address"], data["equipment_class"], data["main_fuse_size"], data["grid_operator"], data["site_description"],
            data["holder_name"], data["holder_address"], data["holder_contact"], data["operator_person"],
            data["builder_name"], data["builder_address"], data["builder_contact"], data["electrical_manager"],
            data["commissioning_inspection"], data["previous_periodic_inspection"], data["this_inspection"], data["next_periodic_inspection"], data["participants"],
            data["inspection_area"], data["inspection_scope_note"], data["inspected_part"], data["method_text"], data["norm_documents"],
            data["decision_summary"], int(data["immediate_danger_flag"]), int(data["reinspection_required_flag"]), data["reinspection_due"],
            data["recommendations_text"], data["attachments_text"], data["distribution_text"], data["signature_name"], data["signature_title"],
        ),
    )
    conn.commit()
    return cursor.lastrowid


def add_finding(report_id: int, severity_class: str, area: str, finding_text: str,
                standard_reference: str, legal_reference: str, corrective_action: str,
                photos: List, photo_caption: str, conn: sqlite3.Connection) -> int:
    next_no = conn.execute(
        "SELECT COALESCE(MAX(finding_no), 0) + 1 FROM findings WHERE report_id = ? AND severity_class = ?",
        (report_id, severity_class),
    ).fetchone()[0]

    cursor = conn.execute(
        """
        INSERT INTO findings (report_id, severity_class, finding_no, area, finding_text, standard_reference, legal_reference, corrective_action)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (report_id, severity_class, next_no, area, finding_text, standard_reference, legal_reference, corrective_action),
    )
    finding_id = cursor.lastrowid
    conn.commit()

    for photo in photos or []:
        save_uploaded_photo(photo, finding_id, photo_caption, conn)

    return finding_id


def list_reports(conn: sqlite3.Connection) -> pd.DataFrame:
    return pd.read_sql_query(
        "SELECT id, report_type, report_number, report_date, site_name, site_address FROM reports ORDER BY id DESC",
        conn,
    )


def get_report(report_id: int, conn: sqlite3.Connection):
    return conn.execute("SELECT * FROM reports WHERE id = ?", (report_id,)).fetchone()


def get_findings(report_id: int, conn: sqlite3.Connection):
    return conn.execute(
        "SELECT * FROM findings WHERE report_id = ? ORDER BY severity_class, finding_no",
        (report_id,),
    ).fetchall()


def get_finding_photos(finding_id: int, conn: sqlite3.Connection):
    return conn.execute(
        "SELECT * FROM finding_photos WHERE finding_id = ? ORDER BY id",
        (finding_id,),
    ).fetchall()


def add_measurement(report_id: int, row: dict, conn: sqlite3.Connection) -> None:
    conn.execute(
        """
        INSERT INTO measurements (
            report_id, center_name, group_name, ik_a, zk_ohm, protection_a,
            phase_order, voltage_v, idn_ma, trip_time_ms, delta_current_ma
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            report_id,
            row.get("center_name", ""), row.get("group_name", ""), row.get("ik_a", ""), row.get("zk_ohm", ""), row.get("protection_a", ""),
            row.get("phase_order", ""), row.get("voltage_v", ""), row.get("idn_ma", ""), row.get("trip_time_ms", ""), row.get("delta_current_ma", ""),
        ),
    )
    conn.commit()


def get_measurements(report_id: int, conn: sqlite3.Connection) -> pd.DataFrame:
    return pd.read_sql_query("SELECT * FROM measurements WHERE report_id = ?", conn, params=(report_id,))


def render_report_html(report, findings, conn: sqlite3.Connection) -> str:
    severity_titles = {
        "0": "5. Välittömän vaaran aiheuttaneet puutteet ja viat",
        "1": "6. Vakavat puutteet ja viat",
        "2": "7. Puutteet ja viat, jotka vähentävät käytön turvallisuutta",
    }

    findings_html = ""
    for sev in ["0", "1", "2"]:
        items = [f for f in findings if f["severity_class"] == sev]
        rows = ""
        for f in items:
            refs = []
            if f["standard_reference"]:
                refs.append(f"Standardi: {f['standard_reference']}")
            if f["legal_reference"]:
                refs.append(f"Laki/asetus: {f['legal_reference']}")
            ref_text = "<br>".join(refs)

            photos = get_finding_photos(f["id"], conn)
            photo_html = ""
            for p in photos:
                photo_html += (
                    f"<div style='margin-top:8px;'>"
                    f"<img src='file://{p['file_path']}' style='max-width:420px; border:1px solid #ccc; padding:4px;'/>"
                    f"<div style='font-size:12px; color:#555;'>{p['caption'] or ''}</div>"
                    f"</div>"
                )

            rows += f"""
            <div style='border:1px solid #ddd; padding:10px; margin:10px 0; border-radius:8px;'>
                <div><strong>Nro {f['finding_no']}</strong> | Alue: {f['area'] or '-'}</div>
                <div style='margin-top:6px;'><strong>Puutteen kuvaus:</strong> {f['finding_text']}</div>
                <div style='margin-top:6px;'>{ref_text}</div>
                <div style='margin-top:6px;'><strong>Korjaus / tekijä:</strong> {f['corrective_action'] or '-'}</div>
                {photo_html}
            </div>
            """
        findings_html += f"<h3>{severity_titles[sev]}</h3>{rows or '<p>Ei kirjattuja puutteita.</p>'}"

    html = f"""
    <html>
    <head>
        <meta charset='utf-8'>
        <title>Tarkastuspöytäkirja {report['report_number']}</title>
    </head>
    <body style='font-family:Arial,sans-serif; max-width:960px; margin:0 auto; padding:24px;'>
        <h1>Sähkölaitteiston tarkastuspöytäkirja</h1>
        <h2>{report['report_type'].capitalize()} | Nro {report['report_number']}</h2>

        <h3>1. Tarkastettu sähkölaitteisto</h3>
        <p><strong>Kohde:</strong> {report['site_name'] or '-'}<br>
        <strong>Osoite:</strong> {report['site_address'] or '-'}<br>
        <strong>Luokka:</strong> {report['equipment_class'] or '-'}<br>
        <strong>Ylivirtasuoja / liittymä:</strong> {report['main_fuse_size'] or '-'}<br>
        <strong>Jakeluverkon haltija:</strong> {report['grid_operator'] or '-'}<br>
        <strong>Kuvaus:</strong> {report['site_description'] or '-'}</p>

        <h3>2. Tarkastuksen kuvaus ja laajuus</h3>
        <p><strong>Tarkastuspaikka:</strong> {report['place'] or '-'}<br>
        <strong>Tarkastuspäivä:</strong> {report['report_date']}<br>
        <strong>Tarkastettu alue:</strong> {report['inspection_area'] or '-'}<br>
        <strong>Tarkennus:</strong> {report['inspection_scope_note'] or '-'}<br>
        <strong>Tarkastettu osa:</strong> {report['inspected_part'] or '-'}<br>
        <strong>Menetelmä:</strong> {report['method_text'] or '-'}<br>
        <strong>Normiasiakirjat:</strong> {report['norm_documents'] or '-'}</p>

        <h3>3. Päätös</h3>
        <p>{report['decision_summary'] or '-'}</p>

        {findings_html}

        <h3>8. Suositukset</h3>
        <p>{report['recommendations_text'] or '-'}</p>

        <h3>Liitteet ja jakelu</h3>
        <p><strong>Liitteet:</strong> {report['attachments_text'] or '-'}<br>
        <strong>Jakelu:</strong> {report['distribution_text'] or '-'}</p>

        <h3>Allekirjoitus</h3>
        <p>{report['place'] or '-'} {report['report_date']}<br>
        {report['signature_name'] or '-'}<br>
        {report['signature_title'] or '-'}</p>
    </body>
    </html>
    """
    return html


def _load_font_for_pdf() -> str:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            try:
                pdfmetrics.registerFont(TTFont("AppFont", candidate))
                return "AppFont"
            except Exception:
                pass
    return "Helvetica"


def _pdf_write_wrapped(c: canvas.Canvas, text: str, x: float, y: float, max_width: float, line_height: float, font_name: str, font_size: int) -> float:
    words = str(text or "").split()
    line = ""
    c.setFont(font_name, font_size)
    for word in words:
        test = f"{line} {word}".strip()
        if c.stringWidth(test, font_name, font_size) <= max_width:
            line = test
        else:
            c.drawString(x, y, line)
            y -= line_height
            line = word
            if y < 20 * mm:
                c.showPage()
                c.setFont(font_name, font_size)
                y = 270 * mm
    if line:
        c.drawString(x, y, line)
        y -= line_height
    return y


def build_pdf_bytes(report, findings, conn: sqlite3.Connection) -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    font_name = _load_font_for_pdf()
    width, height = A4
    x = 15 * mm
    y = height - 15 * mm

    c.setTitle(f"Tarkastuspoytakirja_{report['report_number']}")
    c.setFont(font_name, 16)
    c.drawString(x, y, "Sahkolaitteiston tarkastuspoytakirja")
    y -= 10 * mm
    c.setFont(font_name, 12)
    y = _pdf_write_wrapped(c, f"Tyyppi: {report['report_type']} | Nro: {report['report_number']}", x, y, width - 30 * mm, 5 * mm, font_name, 12)
    y = _pdf_write_wrapped(c, f"Kohde: {report['site_name'] or '-'}", x, y, width - 30 * mm, 5 * mm, font_name, 11)
    y = _pdf_write_wrapped(c, f"Osoite: {report['site_address'] or '-'}", x, y, width - 30 * mm, 5 * mm, font_name, 11)
    y = _pdf_write_wrapped(c, f"Paikka ja pvm: {report['place'] or '-'} {report['report_date']}", x, y, width - 30 * mm, 5 * mm, font_name, 11)
    y -= 4 * mm

    sections = [
        ("Tarkastuksen kuvaus", report["method_text"]),
        ("Normiasiakirjat", report["norm_documents"]),
        ("Paatos", report["decision_summary"]),
        ("Suositukset", report["recommendations_text"]),
    ]
    for title, content in sections:
        c.setFont(font_name, 13)
        c.drawString(x, y, title)
        y -= 6 * mm
        y = _pdf_write_wrapped(c, content or "-", x, y, width - 30 * mm, 5 * mm, font_name, 11)
        y -= 2 * mm

    severity_titles = {
        "0": "Valittoman vaaran puutteet",
        "1": "Vakavat puutteet",
        "2": "Turvallisuutta vahentavat puutteet",
    }
    for sev in ["0", "1", "2"]:
        group = [f for f in findings if f["severity_class"] == sev]
        c.setFont(font_name, 13)
        if y < 40 * mm:
            c.showPage()
            y = height - 15 * mm
        c.drawString(x, y, severity_titles[sev])
        y -= 6 * mm
        if not group:
            y = _pdf_write_wrapped(c, "Ei kirjauksia.", x, y, width - 30 * mm, 5 * mm, font_name, 11)
            continue
        for frow in group:
            text = f"#{frow['finding_no']} | {frow['area'] or '-'} | {frow['finding_text']}"
            refs = f"Standardi: {frow['standard_reference'] or '-'} | Laki: {frow['legal_reference'] or '-'}"
            y = _pdf_write_wrapped(c, text, x, y, width - 30 * mm, 5 * mm, font_name, 11)
            y = _pdf_write_wrapped(c, refs, x, y, width - 30 * mm, 5 * mm, font_name, 10)
            y = _pdf_write_wrapped(c, f"Korjaus / tekija: {frow['corrective_action'] or '-'}", x, y, width - 30 * mm, 5 * mm, font_name, 10)
            y -= 2 * mm
            if y < 25 * mm:
                c.showPage()
                y = height - 15 * mm

    c.save()
    buffer.seek(0)
    return buffer.getvalue()


def build_docx_bytes(report, findings, conn: sqlite3.Connection) -> bytes:
    doc = Document()
    doc.add_heading("Sähkölaitteiston tarkastuspöytäkirja", 0)
    doc.add_paragraph(f"{report['report_type'].capitalize()} | Nro {report['report_number']}")

    doc.add_heading("1. Tarkastettu sähkölaitteisto", level=1)
    fields = [
        ("Kohde", report["site_name"]),
        ("Osoite", report["site_address"]),
        ("Luokka", report["equipment_class"]),
        ("Ylivirtasuoja / liittymä", report["main_fuse_size"]),
        ("Jakeluverkon haltija", report["grid_operator"]),
        ("Kuvaus", report["site_description"]),
    ]
    for label, value in fields:
        doc.add_paragraph(f"{label}: {value or '-'}")

    doc.add_heading("2. Tarkastuksen kuvaus ja laajuus", level=1)
    doc.add_paragraph(f"Paikka: {report['place'] or '-'}")
    doc.add_paragraph(f"Päiväys: {report['report_date']}")
    doc.add_paragraph(f"Menetelmä: {report['method_text'] or '-'}")
    doc.add_paragraph(f"Normiasiakirjat: {report['norm_documents'] or '-'}")

    doc.add_heading("3. Päätös", level=1)
    doc.add_paragraph(report["decision_summary"] or "-")

    severity_titles = {
        "0": "5. Välittömän vaaran aiheuttaneet puutteet ja viat",
        "1": "6. Vakavat puutteet ja viat",
        "2": "7. Puutteet ja viat, jotka vähentävät käytön turvallisuutta",
    }
    for sev in ["0", "1", "2"]:
        doc.add_heading(severity_titles[sev], level=1)
        group = [f for f in findings if f["severity_class"] == sev]
        if not group:
            doc.add_paragraph("Ei kirjauksia.")
        for frow in group:
            doc.add_paragraph(f"Nro {frow['finding_no']} | Alue: {frow['area'] or '-'}", style="List Bullet")
            doc.add_paragraph(f"Puutteen kuvaus: {frow['finding_text']}")
            doc.add_paragraph(f"Standardi: {frow['standard_reference'] or '-'}")
            doc.add_paragraph(f"Laki/asetus: {frow['legal_reference'] or '-'}")
            doc.add_paragraph(f"Korjaus / tekijä: {frow['corrective_action'] or '-'}")

            for photo in get_finding_photos(frow["id"], conn):
                p = Path(photo["file_path"])
                if p.exists():
                    try:
                        doc.add_picture(str(p), width=Inches(4.8))
                        if photo["caption"]:
                            doc.add_paragraph(photo["caption"])
                    except Exception:
                        doc.add_paragraph(f"[Kuvaa ei voitu lisätä: {photo['original_filename']}]")

    doc.add_heading("8. Suositukset", level=1)
    doc.add_paragraph(report["recommendations_text"] or "-")
    doc.add_heading("Liitteet ja jakelu", level=1)
    doc.add_paragraph(f"Liitteet: {report['attachments_text'] or '-'}")
    doc.add_paragraph(f"Jakelu: {report['distribution_text'] or '-'}")
    doc.add_heading("Allekirjoitus", level=1)
    doc.add_paragraph(f"{report['place'] or '-'} {report['report_date']}")
    doc.add_paragraph(report["signature_name"] or "-")
    doc.add_paragraph(report["signature_title"] or "-")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def get_draft_defaults() -> Dict[str, Any]:
    if DRAFT_FILE.exists():
        try:
            data = json.loads(DRAFT_FILE.read_text(encoding="utf-8"))
            return data
        except Exception:
            pass
    return {
        "report_type": "varmennus",
        "report_number": f"xx/{date.today().year}",
        "report_date": str(date.today()),
        "place": "",
        "site_name": "",
        "site_address": "",
        "equipment_class": "",
        "main_fuse_size": "",
        "grid_operator": "",
        "site_description": "",
        "holder_name": "",
        "holder_address": "",
        "holder_contact": "",
        "operator_person": "",
        "builder_name": "",
        "builder_address": "",
        "builder_contact": "",
        "electrical_manager": "",
        "commissioning_inspection": "",
        "previous_periodic_inspection": "",
        "this_inspection": "",
        "next_periodic_inspection": "",
        "participants": "",
        "inspection_area": "",
        "inspection_scope_note": "",
        "inspected_part": "",
        "method_text": DEFAULT_METHOD_TEXT,
        "norm_documents": DEFAULT_NORMS,
        "decision_summary": "",
        "immediate_danger_flag": False,
        "reinspection_required_flag": False,
        "reinspection_due": "",
        "recommendations_text": "",
        "attachments_text": "",
        "distribution_text": "",
        "signature_name": "Riku Leimola",
        "signature_title": "Valtuutettu tarkastaja",
    }


def save_draft_json(data: Dict[str, Any]) -> None:
    payload = dict(data)
    payload["saved_at"] = datetime.now().isoformat(timespec="seconds")
    DRAFT_FILE.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    st.set_page_config(page_title="Tarkastuspöytäkirja", layout="wide")
    st.title("Tarkastuspöytäkirjasovellus")
    st.caption("Streamlit + SQLite + JSON-välitallennus + PDF/DOCX-vienti + AI-viittausehdotus")

    init_db()
    conn = get_conn()
    draft_defaults = get_draft_defaults()

    tab1, tab2, tab3 = st.tabs(["1. Uusi pöytäkirja", "2. Lisää puutteet ja kuvat", "3. Esikatselu ja vienti"])

    with tab1:
        st.subheader("Luo uusi tarkastuspöytäkirja")
        with st.form("create_report_form"):
            col1, col2 = st.columns(2)
            with col1:
                report_type = st.selectbox("Pöytäkirjatyyppi", ["varmennus", "määräaikainen", "vaatimustenmukaisuus"], index=["varmennus", "määräaikainen", "vaatimustenmukaisuus"].index(draft_defaults.get("report_type", "varmennus")))
                report_number = st.text_input("Pöytäkirjan numero", value=draft_defaults.get("report_number", f"xx/{date.today().year}"))
                draft_date = draft_defaults.get("report_date", str(date.today()))
                report_date = st.date_input("Päivämäärä", value=date.fromisoformat(draft_date))
                place = st.text_input("Paikkakunta", value=draft_defaults.get("place", ""))

                site_name = st.text_input("Kohteen nimi", value=draft_defaults.get("site_name", ""))
                site_address = st.text_input("Sijaintiosoite", value=draft_defaults.get("site_address", ""))
                equipment_class = st.text_input("Sähkölaitteiston luokka", value=draft_defaults.get("equipment_class", ""))
                main_fuse_size = st.text_input("Ylivirtasuoja / liittymän koko", value=draft_defaults.get("main_fuse_size", ""))
                grid_operator = st.text_input("Jakeluverkon haltija", value=draft_defaults.get("grid_operator", ""))
                site_description = st.text_area("Kuvaus / tarkennus", value=draft_defaults.get("site_description", ""))

                holder_name = st.text_input("Haltija", value=draft_defaults.get("holder_name", ""))
                holder_address = st.text_input("Haltijan osoite", value=draft_defaults.get("holder_address", ""))
                holder_contact = st.text_input("Yhteyshenkilö", value=draft_defaults.get("holder_contact", ""))
                operator_person = st.text_input("KJ / käytöstä vastaava", value=draft_defaults.get("operator_person", ""))

            with col2:
                builder_name = st.text_input("Rakentaja", value=draft_defaults.get("builder_name", ""))
                builder_address = st.text_input("Rakentajan osoite", value=draft_defaults.get("builder_address", ""))
                builder_contact = st.text_input("Rakentajan yhteyshenkilö", value=draft_defaults.get("builder_contact", ""))
                electrical_manager = st.text_input("Sähkötöiden johtaja", value=draft_defaults.get("electrical_manager", ""))

                commissioning_inspection = st.text_input("Käyttöönottotarkastus", value=draft_defaults.get("commissioning_inspection", ""))
                previous_periodic_inspection = st.text_input("Edellinen määräaikaistarkastus", value=draft_defaults.get("previous_periodic_inspection", ""))
                this_inspection = st.text_input("Tämä tarkastus", value=draft_defaults.get("this_inspection", ""))
                next_periodic_inspection = st.text_input("Seuraava määräaikaistarkastus", value=draft_defaults.get("next_periodic_inspection", ""))
                participants = st.text_area("Tarkastuksessa mukana", value=draft_defaults.get("participants", ""))

                inspection_area = st.text_input("Tarkastettu alue", value=draft_defaults.get("inspection_area", ""))
                inspection_scope_note = st.text_input("Tarkenne tarvittaessa", value=draft_defaults.get("inspection_scope_note", ""))
                inspected_part = st.text_input("Tarkastettu laitteisto tai sen osa", value=draft_defaults.get("inspected_part", ""))
                method_text = st.text_area("Tarkastusmenetelmä", value=draft_defaults.get("method_text", DEFAULT_METHOD_TEXT), height=120)
                norm_documents = st.text_area("Normiasiakirjat", value=draft_defaults.get("norm_documents", DEFAULT_NORMS))

            st.markdown("### Päätös ja allekirjoitus")
            decision_summary = st.text_area("Päätös / yhteenveto", value=draft_defaults.get("decision_summary", ""))
            c1, c2 = st.columns(2)
            with c1:
                immediate_danger_flag = st.checkbox("Välitön vaara todettu", value=bool(draft_defaults.get("immediate_danger_flag", False)))
                reinspection_required_flag = st.checkbox("Uusintatarkastus vaaditaan", value=bool(draft_defaults.get("reinspection_required_flag", False)))
            with c2:
                reinspection_due = st.text_input("Uusintatarkastus sovittu pidettäväksi", value=draft_defaults.get("reinspection_due", ""))

            recommendations_text = st.text_area("Suositukset", value=draft_defaults.get("recommendations_text", ""))
            attachments_text = st.text_area("Liitteet", value=draft_defaults.get("attachments_text", ""))
            distribution_text = st.text_area("Jakelu", value=draft_defaults.get("distribution_text", ""))
            signature_name = st.text_input("Allekirjoittaja", value=draft_defaults.get("signature_name", "Riku Leimola"))
            signature_title = st.text_input("Titteli", value=draft_defaults.get("signature_title", "Valtuutettu tarkastaja"))

            save_col, create_col, clear_col = st.columns(3)
            save_draft_btn = save_col.form_submit_button("Tallenna väliluonnos JSON")
            submitted = create_col.form_submit_button("Tallenna pöytäkirja", type="primary")
            clear_draft_btn = clear_col.form_submit_button("Tyhjennä luonnos")

            draft_payload = {
                "report_type": report_type,
                "report_number": report_number,
                "report_date": str(report_date),
                "place": place,
                "site_name": site_name,
                "site_address": site_address,
                "equipment_class": equipment_class,
                "main_fuse_size": main_fuse_size,
                "grid_operator": grid_operator,
                "site_description": site_description,
                "holder_name": holder_name,
                "holder_address": holder_address,
                "holder_contact": holder_contact,
                "operator_person": operator_person,
                "builder_name": builder_name,
                "builder_address": builder_address,
                "builder_contact": builder_contact,
                "electrical_manager": electrical_manager,
                "commissioning_inspection": commissioning_inspection,
                "previous_periodic_inspection": previous_periodic_inspection,
                "this_inspection": this_inspection,
                "next_periodic_inspection": next_periodic_inspection,
                "participants": participants,
                "inspection_area": inspection_area,
                "inspection_scope_note": inspection_scope_note,
                "inspected_part": inspected_part,
                "method_text": method_text,
                "norm_documents": norm_documents,
                "decision_summary": decision_summary,
                "immediate_danger_flag": immediate_danger_flag,
                "reinspection_required_flag": reinspection_required_flag,
                "reinspection_due": reinspection_due,
                "recommendations_text": recommendations_text,
                "attachments_text": attachments_text,
                "distribution_text": distribution_text,
                "signature_name": signature_name,
                "signature_title": signature_title,
            }

            if save_draft_btn:
                save_draft_json(draft_payload)
                st.success(f"Väliluonnos tallennettu: {DRAFT_FILE}")

            if clear_draft_btn:
                if DRAFT_FILE.exists():
                    DRAFT_FILE.unlink()
                st.success("Luonnos poistettu. Päivitä sivu.")

            if submitted:
                report_id = create_report(draft_payload, conn)
                save_draft_json(draft_payload)
                st.success(f"Pöytäkirja tallennettu. Report ID: {report_id}")

        if DRAFT_FILE.exists():
            st.download_button(
                "Lataa luonnos JSON",
                data=DRAFT_FILE.read_bytes(),
                file_name=DRAFT_FILE.name,
                mime="application/json",
            )

    with tab2:
        st.subheader("Lisää puutteet ja kuvat")
        reports_df = list_reports(conn)
        if reports_df.empty:
            st.info("Luo ensin pöytäkirja välilehdellä 1.")
        else:
            report_options = {f"{r['id']} | {r['report_type']} | {r['report_number']} | {r['site_name']}": r['id'] for _, r in reports_df.iterrows()}
            selected_label = st.selectbox("Valitse pöytäkirja", list(report_options.keys()))
            selected_report_id = report_options[selected_label]

            if "ai_selected_standard" not in st.session_state:
                st.session_state.ai_selected_standard = ""
            if "ai_selected_legal" not in st.session_state:
                st.session_state.ai_selected_legal = ""

            finding_text_input = st.text_area("Puutteen kuvaus AI-ehdotusta varten", height=120, key="finding_text_ai")
            ai_col1, ai_col2 = st.columns([1, 2])
            with ai_col1:
                ai_btn = st.button("AI ehdota viitteet")
            if ai_btn:
                st.session_state.ai_suggestions = suggest_references_from_csv(finding_text_input, limit=5)

            suggestions = st.session_state.get("ai_suggestions", []) if finding_text_input.strip() else []
            if suggestions:
                st.markdown("#### AI-viittausehdotukset")
                for idx, s in enumerate(suggestions, start=1):
                    with st.container(border=True):
                        st.write(f"**{idx}. {s['label']}**")
                        if s["heading"]:
                            st.write(f"Otsikko: {s['heading']}")
                        if s["source_type"]:
                            st.caption(f"Lähdetyyppi: {s['source_type']}")
                        st.write(s["excerpt"])
                        apply_col1, apply_col2 = st.columns(2)
                        if apply_col1.button(f"Käytä standardiksi #{idx}", key=f"use_std_{idx}"):
                            st.session_state.ai_selected_standard = s["label"]
                        if apply_col2.button(f"Käytä lakiviitteeksi #{idx}", key=f"use_law_{idx}"):
                            st.session_state.ai_selected_legal = s["label"]
            elif ai_btn:
                st.info("CSV-tietopohjasta ei löytynyt osumia.")

            with st.form("add_finding_form"):
                c1, c2 = st.columns(2)
                with c1:
                    severity_class = st.selectbox("Puutteluokka", ["0", "1", "2"], format_func=lambda x: {"0": "0 - välitön vaara", "1": "1 - vakava puute", "2": "2 - turvallisuutta vähentävä puute"}[x])
                    area = st.text_input("Alue", placeholder="esim. Pääkeskus, ryhmäkeskus, lääkintätila")
                    standard_reference = st.text_input("Standardiviite", value=st.session_state.get("ai_selected_standard", ""), placeholder="esim. SFS 6000:2022 kohta 514.5.1")
                    legal_reference = st.text_input("Laki-/asetusviite", value=st.session_state.get("ai_selected_legal", ""), placeholder="esim. STL 1135/2016 45 §")
                with c2:
                    corrective_action = st.text_input("Korjattu / tekijä")
                    photo_caption = st.text_input("Kuvateksti", placeholder="esim. Keskuksen puutteellinen merkintä")
                    photos = st.file_uploader(
                        "Lisää 1 tai useampi valokuva",
                        type=["jpg", "jpeg", "png", "webp"],
                        accept_multiple_files=True,
                    )

                finding_text = st.text_area("Puutteen kuvaus", value=finding_text_input, height=140)
                add_finding_btn = st.form_submit_button("Tallenna puutekohta", type="primary")
                if add_finding_btn:
                    if not finding_text.strip():
                        st.warning("Kirjoita puutteen kuvaus.")
                    else:
                        finding_id = add_finding(
                            selected_report_id, severity_class, area, finding_text,
                            standard_reference, legal_reference, corrective_action,
                            photos, photo_caption, conn,
                        )
                        st.success(f"Puutekohta tallennettu. Finding ID: {finding_id}")
                        st.session_state.ai_selected_standard = ""
                        st.session_state.ai_selected_legal = ""
                        st.session_state.ai_suggestions = []

            st.markdown("### Tallennetut puutteet")
            findings = get_findings(selected_report_id, conn)
            if not findings:
                st.info("Tähän pöytäkirjaan ei ole vielä lisätty puutteita.")
            for f in findings:
                with st.container(border=True):
                    st.markdown(f"**Luokka {f['severity_class']} / nro {f['finding_no']}** — {f['area'] or '-'}")
                    st.write(f["finding_text"])
                    if f["standard_reference"]:
                        st.caption(f"Standardi: {f['standard_reference']}")
                    if f["legal_reference"]:
                        st.caption(f"Laki/asetus: {f['legal_reference']}")
                    if f["corrective_action"]:
                        st.caption(f"Korjattu / tekijä: {f['corrective_action']}")

                    photo_rows = get_finding_photos(f["id"], conn)
                    if photo_rows:
                        cols = st.columns(min(3, len(photo_rows)))
                        for idx, p in enumerate(photo_rows):
                            with cols[idx % len(cols)]:
                                st.image(p["file_path"], caption=p["caption"] or p["original_filename"], use_container_width=True)

    with tab3:
        st.subheader("Esikatselu ja vienti")
        reports_df = list_reports(conn)
        if reports_df.empty:
            st.info("Ei vielä tallennettuja pöytäkirjoja.")
        else:
            report_options = {f"{r['id']} | {r['report_type']} | {r['report_number']} | {r['site_name']}": r['id'] for _, r in reports_df.iterrows()}
            selected_label = st.selectbox("Valitse pöytäkirja esikatseluun", list(report_options.keys()), key="preview_report")
            selected_report_id = report_options[selected_label]

            report = get_report(selected_report_id, conn)
            findings = get_findings(selected_report_id, conn)
            measurements_df = get_measurements(selected_report_id, conn)

            st.markdown(f"## {report['report_type'].capitalize()} — {report['report_number']}")
            st.write(f"**Kohde:** {report['site_name'] or '-'}")
            st.write(f"**Osoite:** {report['site_address'] or '-'}")
            st.write(f"**Päiväys:** {report['report_date']}")
            st.write(f"**Päätös:** {report['decision_summary'] or '-'}")

            for sev, title in [("0", "Välittömän vaaran puutteet"), ("1", "Vakavat puutteet"), ("2", "Turvallisuutta vähentävät puutteet")]:
                st.markdown(f"### {title}")
                sev_findings = [f for f in findings if f["severity_class"] == sev]
                if not sev_findings:
                    st.write("Ei kirjauksia")
                    continue
                for f in sev_findings:
                    with st.container(border=True):
                        st.write(f"**#{f['finding_no']}** | {f['area'] or '-'}")
                        st.write(f["finding_text"])
                        if f["standard_reference"]:
                            st.write(f"Standardi: {f['standard_reference']}")
                        if f["legal_reference"]:
                            st.write(f"Laki/asetus: {f['legal_reference']}")
                        photos = get_finding_photos(f["id"], conn)
                        if photos:
                            photo_cols = st.columns(min(3, len(photos)))
                            for idx, p in enumerate(photos):
                                with photo_cols[idx % len(photo_cols)]:
                                    st.image(p["file_path"], caption=p["caption"] or p["original_filename"], use_container_width=True)

            if not measurements_df.empty:
                st.markdown("### Mittaustulokset")
                st.dataframe(measurements_df, use_container_width=True)

            html = render_report_html(report, findings, conn)
            pdf_bytes = build_pdf_bytes(report, findings, conn)
            docx_bytes = build_docx_bytes(report, findings, conn)

            c1, c2, c3 = st.columns(3)
            c1.download_button(
                "Lataa HTML-pöytäkirja",
                data=html,
                file_name=f"tarkastuspoytakirja_{report['report_number'].replace('/', '_')}.html",
                mime="text/html",
            )
            c2.download_button(
                "Lataa PDF-pöytäkirja",
                data=pdf_bytes,
                file_name=f"tarkastuspoytakirja_{report['report_number'].replace('/', '_')}.pdf",
                mime="application/pdf",
            )
            c3.download_button(
                "Lataa Word-pöytäkirja",
                data=docx_bytes,
                file_name=f"tarkastuspoytakirja_{report['report_number'].replace('/', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    conn.close()


if __name__ == "__main__":
    main()
